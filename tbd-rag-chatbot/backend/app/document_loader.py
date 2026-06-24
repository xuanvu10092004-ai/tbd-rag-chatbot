# -*- coding: utf-8 -*-
import logging
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup

from app.config import settings
from app.seed_urls import ALLOWED_DOMAINS

logger = logging.getLogger(__name__)

# Các định dạng file được hỗ trợ
SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".md", ".markdown"}


def _make_document(content: str, source: str, source_type: str, title: str = "") -> dict:
    """
    Tao dict chuan cho mot tai lieu da tai.
    Day la cau truc du lieu truoc khi tach chunk.
    """
    return {
        "content": content.strip(),
        "source": source,
        "source_type": source_type,
        "title": title or source,
    }


def _clean_text(text: str) -> str:
    """
    Lam sach van ban: xoa khoang trang thua, chuyen hoa ky tu xuong dong.
    Giu lai cau truc doan van de RecursiveCharacterTextSplitter hoat dong tot hon.
    """
    import re
    # Gop nhieu dong trong thanh mot dong trong
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Xoa khoang trang dau/cuoi moi dong
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(lines)
    # Xoa khoang trang dau/cuoi toan bo van ban
    return text.strip()


def load_txt(path: str) -> list[dict]:
    """
    Tai file van ban thuan (TXT hoac Markdown).
    Tra ve danh sach chua mot tai lieu.

    Args:
        path: Duong dan tuyet doi den file

    Returns:
        Danh sach dict tai lieu, hoac list rong neu that bai
    """
    try:
        content = Path(path).read_text(encoding="utf-8")
        content = _clean_text(content)
        if not content:
            logger.warning("File TXT rong: %s", path)
            return []
        logger.info("Da tai file TXT: %s (%d ky tu)", path, len(content))
        return [_make_document(content, path, "file", Path(path).name)]
    except Exception as e:
        logger.error("Loi khi tai file TXT '%s': %s", path, str(e))
        return []


def load_markdown(path: str) -> list[dict]:
    """
    Tai file Markdown - xu ly tuong tu TXT vi ChromaDB luu van ban thuan.
    """
    return load_txt(path)


def load_pdf(path: str) -> list[dict]:
    """
    Tai file PDF su dung thu vien pypdf.
    Gop noi dung cac trang thanh mot tai lieu.

    Args:
        path: Duong dan tuyet doi den file PDF

    Returns:
        Danh sach dict tai lieu, hoac list rong neu that bai
    """
    try:
        # pyrefly: ignore [missing-import]
        from pypdf import PdfReader

        reader = PdfReader(path)
        pages_text = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages_text.append(text)

        if not pages_text:
            logger.warning("File PDF khong co noi dung trich xuat duoc: %s", path)
            return []

        content = _clean_text("\n\n".join(pages_text))
        logger.info("Da tai file PDF: %s (%d trang, %d ky tu)", path, len(reader.pages), len(content))
        return [_make_document(content, path, "file", Path(path).stem)]

    except ImportError:
        logger.error("Chua cai dat pypdf. Chay: pip install pypdf")
        return []
    except Exception as e:
        logger.error("Loi khi tai file PDF '%s': %s", path, str(e))
        return []


def load_docx(path: str) -> list[dict]:
    """
    Tai file DOCX su dung thu vien python-docx.
    Lay noi dung cac doan van (paragraphs) va bang (tables).

    Args:
        path: Duong dan tuyet doi den file DOCX

    Returns:
        Danh sach dict tai lieu, hoac list rong neu that bai
    """
    try:
        from docx import Document

        doc = Document(path)
        parts = []

        # Lay noi dung cac doan van
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Lay noi dung cac bang
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        if not parts:
            logger.warning("File DOCX khong co noi dung: %s", path)
            return []

        content = _clean_text("\n\n".join(parts))
        logger.info("Da tai file DOCX: %s (%d ky tu)", path, len(content))
        return [_make_document(content, path, "file", Path(path).stem)]

    except ImportError:
        logger.error("Chua cai dat python-docx. Chay: pip install python-docx")
        return []
    except Exception as e:
        logger.error("Loi khi tai file DOCX '%s': %s", path, str(e))
        return []


def _is_allowed_url(url: str) -> bool:
    """
    Kiểm tra URL có thuộc danh sách tên miền được phép hay không.
    Bảo mật: từ chối tất cả URL ngoài whitelist.
    """
    try:
        parsed = urlparse(url)
        host = parsed.netloc.lower()
        if host.startswith("www."):
            host = host[4:]
        return host in ALLOWED_DOMAINS
    except Exception:
        return False


def _extract_page_title(soup: BeautifulSoup) -> str:
    """
    Trich xuat tieu de trang tu the <title> hoac <h1> dau tien.
    """
    # Thu lay tu <title>
    title_tag = soup.find("title")
    if title_tag and title_tag.get_text(strip=True):
        return title_tag.get_text(strip=True)

    # Thu lay tu <h1>
    h1_tag = soup.find("h1")
    if h1_tag and h1_tag.get_text(strip=True):
        return h1_tag.get_text(strip=True)

    return ""


def _extract_main_content(soup: BeautifulSoup) -> str:
    """
    Trích xuất nội dung chính từ trang web.
    Bỏ qua các phần: header, footer, nav, sidebar, script, style.
    Đồng thời định dạng lại bảng biểu và tiêu đề để giữ cấu trúc thông tin.
    """
    # Xóa các phần không cần thiết
    for tag in soup.find_all(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
        tag.decompose()

    # Thử tìm khu vực nội dung chính
    main_content = (
        soup.find("main")
        or soup.find("article")
        or soup.find(attrs={"class": lambda c: c and any(
            kw in " ".join(c) for kw in ["content", "main", "article", "post", "entry"]
        )})
        or soup.find("body")
    )

    if not main_content:
        main_content = soup

    # Định dạng lại bảng biểu (tables) thành văn bản dễ đọc trước khi trích xuất
    for table in main_content.find_all("table"):
        rows_text = []
        for row in table.find_all("tr"):
            cells = [cell.get_text(strip=True) for cell in row.find_all(["td", "th"])]
            row_str = " | ".join(c for c in cells if c)
            if row_str:
                rows_text.append(row_str)
        if rows_text:
            table_text = "\n" + "\n".join(rows_text) + "\n"
            new_tag = soup.new_tag("div")
            new_tag.string = table_text
            table.replace_with(new_tag)

    # Định dạng lại tiêu đề (headings) để giữ cấu trúc văn bản phân cấp
    for h in main_content.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
        level = int(h.name[1])
        prefix = "#" * level + " "
        h_text = h.get_text(strip=True)
        if h_text:
            new_tag = soup.new_tag("div")
            new_tag.string = f"\n\n{prefix}{h_text}\n\n"
            h.replace_with(new_tag)

    return _clean_text(main_content.get_text(separator="\n"))


def load_url(url: str) -> list[dict]:
    """
    Tai noi dung trang web tu URL chinh thuc cua TBD.
    Chi chap nhan URL tu domain https://tbd.edu.vn/.
    Khong thuc hien crawl de quy - chi lay dung URL duoc cung cap.

    Args:
        url: URL can tai (phai thuoc domain tbd.edu.vn)

    Returns:
        Danh sach dict tai lieu, hoac list rong neu that bai hoac khong hop le
    """
    # Kiem tra domain whitelist
    if not _is_allowed_url(url):
        logger.warning("Từ chối URL ngoài domain cho phép: %s (chỉ cho phép %s)", url, ", ".join(ALLOWED_DOMAINS))
        return []

    try:
        headers = {
            "User-Agent": "TBD-RAG-Bot/1.0 (Educational chatbot for TBD University; contact: admin@tbd.edu.vn)",
            "Accept": "text/html,application/xhtml+xml",
            "Accept-Language": "vi,en;q=0.9",
        }

        response = requests.get(
            url,
            headers=headers,
            timeout=settings.URL_REQUEST_TIMEOUT,
            allow_redirects=True,
        )
        response.raise_for_status()

        # Kiem tra content-type phai la HTML
        content_type = response.headers.get("Content-Type", "")
        if "text/html" not in content_type and "application/xhtml" not in content_type:
            logger.warning("URL khong phai HTML (content-type: %s): %s", content_type, url)
            return []

        soup = BeautifulSoup(response.content, "html.parser")
        title = _extract_page_title(soup)
        content = _extract_main_content(soup)

        if not content or len(content) < 100:
            logger.warning("Noi dung trang qua ngan hoac rong: %s (%d ky tu)", url, len(content))
            return []

        logger.info("Da tai URL: %s | Tieu de: '%s' | %d ky tu", url, title[:50], len(content))
        return [_make_document(content, url, "url", title or url)]

    except requests.exceptions.Timeout:
        logger.error("Timeout khi tai URL (qua %ds): %s", settings.URL_REQUEST_TIMEOUT, url)
        return []
    except requests.exceptions.HTTPError as e:
        logger.error("Loi HTTP khi tai URL '%s': %s", url, str(e))
        return []
    except requests.exceptions.RequestException as e:
        logger.error("Loi ket noi khi tai URL '%s': %s", url, str(e))
        return []
    except Exception as e:
        logger.error("Loi khong xac dinh khi tai URL '%s': %s", url, str(e))
        return []


def load_file(path: str) -> list[dict]:
    """
    Tai file theo dinh dang duoc xac dinh tu extension.
    Tra ve list rong va log WARNING neu dinh dang khong duoc ho tro.

    Args:
        path: Duong dan den file

    Returns:
        Danh sach dict tai lieu
    """
    ext = Path(path).suffix.lower()

    if ext == ".txt":
        return load_txt(path)
    elif ext in {".md", ".markdown"}:
        return load_markdown(path)
    elif ext == ".pdf":
        return load_pdf(path)
    elif ext == ".docx":
        return load_docx(path)
    else:
        logger.warning(
            "Dinh dang file khong duoc ho tro: '%s'. Ho tro: %s",
            ext,
            ", ".join(SUPPORTED_EXTENSIONS),
        )
        return []
